import streamlit as st
import google.generativeai as genai
import os
from docx import Document

# Função para configurar o modelo Gemini com cache
@st.cache_resource
def setup_gemini_model(api_key, system_instruction):
    try:
        genai.configure(api_key=api_key)

        # Configurações de geração
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }

        # Modelo com instruções de sistema
        model = genai.GenerativeModel(
            model_name="gemini-1.5-pro",
            generation_config=generation_config,
            system_instruction=system_instruction
        )

        return model
    except Exception as e:
        st.error(f"Erro ao configurar o modelo: {e}")
        return None

# Função para carregar instruções do sistema
def load_system_instruction(filepath):
    try:
        with open(filepath, "r", encoding="utf-8") as file:
            return file.read()
    except FileNotFoundError:
        st.error(f"Arquivo '{filepath}' não encontrado.")
        return ""
    except Exception as e:
        st.error(f"Erro ao carregar instruções: {e}")
        return ""


def main():
    st.title("Relatório ETP | Governo do Estado do Piauí")

    # Input para a chave API
    api_key = st.sidebar.text_input("Digite sua Chave API do Gemini", type="password")

    # Carregar instruções do sistema
    system_instruction = load_system_instruction(os.path.join(os.path.dirname(__file__), "comandos.txt"))

    # Verifica se a chave API foi fornecida
    if api_key and system_instruction:
        # Configura o modelo Gemini usando cache
        model = setup_gemini_model(api_key, system_instruction)

        if model:
            # Barra lateral para navegação de chats
            st.sidebar.header("Histórico de Chats")
            if 'chat_sessions' not in st.session_state:
                st.session_state.chat_sessions = {}
            if 'active_chat' not in st.session_state:
                st.session_state.active_chat = "Novo Chat"
            if 'report_generated' not in st.session_state:
                st.session_state.report_generated = False
            if 'items_list' not in st.session_state:
                st.session_state.items_list = []

            # Exibe opções de chat na barra lateral
            chat_names = list(st.session_state.chat_sessions.keys()) + ["Novo Chat"]
            selected_chat = st.sidebar.selectbox("Selecione o Chat", chat_names)

            if st.sidebar.button("Criar Novo Chat"):
                chat_name = f"Chat {len(st.session_state.chat_sessions) + 1}"
                st.session_state.chat_sessions[chat_name] = model.start_chat(history=[])
                st.session_state.active_chat = chat_name
                st.session_state.report_generated = False

            if selected_chat != st.session_state.active_chat:
                st.session_state.active_chat = selected_chat

            active_chat = st.session_state.active_chat

            if active_chat not in st.session_state.chat_sessions:
                st.session_state.chat_sessions[active_chat] = model.start_chat(history=[])

            chat_session = st.session_state.chat_sessions[active_chat]

            # Formulário retrátil para adicionar itens
            with st.expander("Formulário para Relatório", expanded=not st.session_state.report_generated):
                # Parte do formulário principal
                with st.form("formulario_relatorio"):
                    setor_requisitante = st.text_input("Setor Requisitante")
                    numero_processo = st.text_input("Número do Processo")
                    informacoes_adicionais = st.text_area("Informações Adicionais")

                    submit_button = st.form_submit_button("Gerar Relatório")

                # Adicionar itens fora do formulário
                st.subheader("Adicionar Itens")
                col1, col2, col3 = st.columns(3)
                with col1:
                    item_nome = st.text_input("Item", key="item_nome")
                with col2:
                    item_quantidade = st.number_input("Quantidade", min_value=1, key="item_quantidade")
                with col3:
                    item_unidade = st.text_input("Unidade de Medida", key="item_unidade")

                if st.button("Adicionar Item"):
                    st.session_state.items_list.append({
                        "nome": item_nome,
                        "quantidade": item_quantidade,
                        "unidade": item_unidade
                    })
                    st.success(f"Item '{item_nome}' adicionado com sucesso!")

                # Mostrar itens adicionados
                if st.session_state.items_list:
                    st.subheader("Itens Adicionados")
                    for idx, item in enumerate(st.session_state.items_list):
                        st.write(f"{idx + 1}. {item['nome']} - {item['quantidade']} {item['unidade']}")

                if submit_button:
                    # Gerar relatório inicial
                    itens_texto = "\n".join(
                        [f"- {item['nome']}: {item['quantidade']} {item['unidade']}" for item in st.session_state.items_list]
                    )
                    initial_message = f"""
                    Setor Requisitante: {setor_requisitante}
                    Número do Processo: {numero_processo}
                    Informações Adicionais: {informacoes_adicionais}
                    Itens:
                    {itens_texto}
                    """
                    try:
                        with st.spinner("Gerando a primeira versão do relatório..."):
                            response = chat_session.send_message(initial_message)
                            st.session_state.report_generated = True

                        # Exibe a resposta do modelo
                        st.success("Relatório gerado com sucesso!")
                        st.chat_message("assistant").write(response.text)

                        # Salvar em Word
                        doc = Document()
                        doc.add_heading(f"Chat: {active_chat}", level=1)
                        doc.add_paragraph(response.text)
                        doc_path = f"{active_chat}.docx"
                        doc.save(doc_path)

                        with open(doc_path, "rb") as file:
                            st.download_button(
                                label="Baixar Documento Word",
                                data=file,
                                file_name=f"{active_chat}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Erro ao gerar o relatório: {e}")

            # Habilita o chat somente após o relatório inicial
            if st.session_state.report_generated:
                st.header(f"Chat Ativo: {active_chat}")
                for message in chat_session.history:
                    if message.role == "user":
                        st.chat_message("user").write(message.parts[0].text)
                    else:
                        st.chat_message("assistant").write(message.parts[0].text)

                # Input de mensagem do usuário
                if prompt := st.chat_input("Digite sua mensagem"):
                    st.chat_message("user").write(prompt)
                    try:
                        with st.spinner("Gerando resposta..."):
                            response = chat_session.send_message(prompt)

                        st.chat_message("assistant").write(response.text)

                    except Exception as e:
                        st.error(f"Erro ao enviar mensagem: {e}")
    else:
        st.warning("Por favor, insira sua chave API do Gemini e certifique-se de que o arquivo './comandos.txt' está disponível.")


# Executa a aplicação
if __name__ == "__main__":
    main()
